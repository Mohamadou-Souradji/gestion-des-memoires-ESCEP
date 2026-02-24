from django.contrib.auth.decorators import login_required
from .models import Filiere, Departement,Classe,AnneeScolaire,Vague,ClotureVagueDepartement
from app_gestion_interne.models import Etudiant, DossierMemoire
from django.shortcuts import render, redirect
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from app_gestion_interne.models import DossierMemoire, Etudiant # Modèles qui sont dans auth
from django.core.exceptions import FieldError
from django.db.models import Q
from django.db import models  # <--- Ajoutez cette ligne


@login_required
def dashboard_directeur(request):
    # --- ACADÉMIQUE ---
    nb_depts = Departement.objects.count()
    nb_etudiants = Etudiant.objects.count()
    # On trie par matricule car 'id' n'existe pas
    derniers_etudiants = Etudiant.objects.select_related('classe').order_by('-matricule')[:5]

    # --- SOUTENANCES & JURYS ---
    nb_soutenances = DossierMemoire.objects.filter(is_soutenu=True).count()
    # On compte les dossiers en attente de validation par le DE
    nb_attente_validation = Soutenance.objects.filter(status='PROPOSE').count()

    # --- SUPERVISION (Contrôle interne) ---
    # On vérifie les dossiers où le pré-dépôt est fait mais pas encore le post-dépôt
    nb_pre_depots = DossierMemoire.objects.filter(is_pre_depot_fait=True, is_post_depot_fait=False).count()

    # --- BIBLIOTHÈQUE ---
    # Mémoires soutenus mais non encore publiés au catalogue
    nb_memoires_attente = DossierMemoire.objects.filter(is_soutenu=True, is_publie=False).count()

    # --- RÉPARTITION PAR DÉPARTEMENT ---
    stats_depts = []
    departements = Departement.objects.all()
    for dept in departements:
        count = Etudiant.objects.filter(classe__filiere__departement=dept).count()
        pourcentage = (count / nb_etudiants * 100) if nb_etudiants > 0 else 0
        stats_depts.append({
            'nom': dept.nom,
            'total_etudiants': count,
            'pourcentage': pourcentage
        })

    context = {
        'nb_depts': nb_depts,
        'nb_etudiants': nb_etudiants,
        'nb_soutenances': nb_soutenances,
        'nb_attente_validation': nb_attente_validation,
        'nb_pre_depots': nb_pre_depots,
        'nb_memoires_attente': nb_memoires_attente,
        'stats_depts': stats_depts,
        'derniers_etudiants': derniers_etudiants,
    }
    return render(request, 'administration/dashboard_de.html', context)

from app_auth.models import User
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.db.models import Q
from app_auth.models import User
from app_administration.models import Departement


# Sécurité : Seul le Directeur des Études peut accéder à cette page
def is_de(user):
    return user.is_authenticated and (user.role == 'DE' or user.is_superuser)


@login_required
@user_passes_test(is_de)
def liste_utilisateurs(request):
    """
    Vue unique pour lister, filtrer, créer et modifier les utilisateurs.
    """
    # --- LOGIQUE DE TRAITEMENT DU FORMULAIRE (POST) ---
    if request.method == "POST":
        user_id = request.POST.get('user_id')
        username = request.POST.get('username')
        email = request.POST.get('email')
        role = request.POST.get('role')
        password = request.POST.get('password')
        dept_id = request.POST.get('departement')
        is_active = request.POST.get('is_active') == 'on'

        # Validation spécifique : Département obligatoire pour Chef de Dept
        if role == 'CHEF_DEPT' and not dept_id:
            messages.error(request, "Le département est obligatoire pour un Chef de Département.")
            return redirect('administration:liste_utilisateurs')

        try:
            if user_id:  # MODIFICATION
                user_obj = get_object_or_404(User, id=user_id)
                user_obj.username = username
                if password:  # On ne change le mot de passe que s'il est saisi
                    user_obj.set_password(password)
                msg = f"Compte {username} mis à jour avec succès."
            else:  # CRÉATION
                if User.objects.filter(username=username).exists():
                    messages.error(request, f"Erreur : Le login '{username}' est déjà utilisé.")
                    return redirect('administration:liste_utilisateurs')

                user_obj = User.objects.create(username=username)
                user_obj.set_password(password if password else "Escep2026")
                msg = f"Utilisateur {username} créé avec succès."

            # Champs communs
            user_obj.email = email
            user_obj.role = role
            user_obj.is_active = is_active

            # On n'assigne le département QUE si le rôle est Chef de Dept
            user_obj.departement_id = dept_id if role == 'CHEF_DEPT' else None

            user_obj.save()
            messages.success(request, msg)

        except Exception as e:
            messages.error(request, f"Une erreur est survenue : {str(e)}")

        return redirect('administration:liste_utilisateurs')

    # --- LOGIQUE D'AFFICHAGE ET FILTRAGE (GET) ---
    query = request.GET.get('q', '')
    role_f = request.GET.get('role_filter', '')
    dept_f = request.GET.get('dept_filter', '')

    # On récupère tous les utilisateurs avec leur département (optimisation SQL)
    utilisateurs = User.objects.select_related('departement').all().order_by('-date_joined')

    if query:
        utilisateurs = utilisateurs.filter(
            Q(username__icontains=query) |
            Q(email__icontains=query) |
            Q(last_name__icontains=query)
        )

    if role_f:
        utilisateurs = utilisateurs.filter(role=role_f)

    if dept_f:
        utilisateurs = utilisateurs.filter(departement_id=dept_f)

    context = {
        'utilisateurs': utilisateurs,
        'roles': User.ROLE_CHOICES,
        'departements': Departement.objects.all(),
    }
    return render(request, 'administration/utilisateurs.html', context)


@login_required
@user_passes_test(is_de)
def utilisateur_delete(request, user_id):
    """
    Vue de suppression sécurisée.
    """
    user_to_del = get_object_or_404(User, id=user_id)

    # Interdire de se supprimer soi-même
    if user_to_del == request.user:
        messages.error(request, "Action interdite : vous ne pouvez pas supprimer votre propre compte administrateur.")
    else:
        nom = user_to_del.username
        user_to_del.delete()
        messages.success(request, f"L'utilisateur {nom} a été supprimé définitivement.")

    return redirect('administration:liste_utilisateurs')


@login_required
def utilisateur_delete(request, user_id):
    # Sécurité : Seul le Directeur des Études ou un Superuser peut supprimer
    if request.user.role != 'DE' and not request.user.is_superuser:
        messages.error(request, "Accès non autorisé.")
        return redirect('administration:dashboard_de')

    user_to_del = get_object_or_404(User, id=user_id)

    # Empêcher l'auto-suppression
    if user_to_del == request.user:
        messages.error(request, "Action impossible : vous ne pouvez pas supprimer votre propre compte.")
    else:
        username = user_to_del.username
        user_to_del.delete()
        messages.success(request, f"L'utilisateur {username} a été supprimé définitivement.")

    return redirect('administration:liste_utilisateurs')

def liste_annees(request):
    annees = AnneeScolaire.objects.all().order_by('-libelle')
    
    if request.method == 'POST':
        libelle = request.POST.get('libelle').strip()
        
        # Verification de l'unicite du libelle
        if AnneeScolaire.objects.filter(libelle=libelle).exists():
            messages.error(request, f"L'annee {libelle} existe deja dans le systeme.")
        else:
            # On cree l'annee (par defaut inactive pour ne pas perturber l'actuelle)
            AnneeScolaire.objects.create(libelle=libelle, est_active=False)
            messages.success(request, f"Annee {libelle} creee avec succes.")
        
        return redirect('administration:liste_annees')

    return render(request, 'administration/liste_annees.html', {'annees': annees})

def activer_annee(request, pk):
    # On desactive toutes les annees d'abord
    AnneeScolaire.objects.update(est_active=False)
    # On active l'annee selectionnee
    annee = get_object_or_404(AnneeScolaire, pk=pk)
    annee.est_active = True
    annee.save()
    messages.success(request, f"L'annee {annee.libelle} est desormais l'annee en cours.")
    return redirect('administration:liste_annees')

def supprimer_annee(request, pk):
    annee = get_object_or_404(AnneeScolaire, pk=pk)
    if annee.est_active:
        messages.error(request, "Impossible de supprimer l'annee en cours.")
    else:
        annee.delete()
        messages.warning(request, "Annee supprimee.")
    return redirect('administration:liste_annees')

#gestion des departements
def liste_departements(request):
    departements = Departement.objects.all().order_by('nom')
    
    if request.method == 'POST':
        nom = request.POST.get('nom').strip()
        
        # Vérification de l'unicité du nom
        if Departement.objects.filter(nom__iexact=nom).exists():
            messages.error(request, f"Le département '{nom}' existe déjà.")
        else:
            Departement.objects.create(nom=nom)
            messages.success(request, f"Département '{nom}' ajouté avec succès.")
        
        return redirect('administration:liste_departements')

    return render(request, 'administration/liste_departements.html', {'departements': departements})

def supprimer_departement(request, pk):
    dept = get_object_or_404(Departement, pk=pk)
    # On peut ajouter une sécurité : ne pas supprimer si des filières y sont liées
    if dept.filieres.exists():
        messages.error(request, "Impossible de supprimer : ce département contient des filières.")
    else:
        dept.delete()
        messages.warning(request, "Département supprimé.")
    return redirect('administration:liste_departements')

def modifier_departement(request, pk):
    dept = get_object_or_404(Departement, pk=pk)
    if request.method == 'POST':
        nouveau_nom = request.POST.get('nom').strip()
        
        # On vérifie si le nom existe déjà sur UN AUTRE département
        if Departement.objects.filter(nom__iexact=nouveau_nom).exclude(pk=pk).exists():
            messages.error(request, f"Le département '{nouveau_nom}' existe déjà.")
        else:
            dept.nom = nouveau_nom
            dept.save()
            messages.success(request, "Département mis à jour avec succès.")
            
    return redirect('administration:liste_departements')

#filieres 


def liste_filieres(request):
    filieres = Filiere.objects.all().select_related('departement').order_by('nom')
    departements = Departement.objects.all().order_by('nom')
    
    if request.method == 'POST':
        nom = request.POST.get('nom').strip()
        code = request.POST.get('code').strip().upper()
        dept_id = request.POST.get('departement')
        
        # Verification de l'unicite (Nom ou Code)
        if Filiere.objects.filter(nom__iexact=nom).exists() or Filiere.objects.filter(code__iexact=code).exists():
            messages.error(request, "Une filiere avec ce nom ou ce code existe deja.")
        else:
            dept = get_object_or_404(Departement, id=dept_id)
            Filiere.objects.create(nom=nom, code=code, departement=dept)
            messages.success(request, f"Filiere {code} creee avec succes.")
        
        return redirect('administration:liste_filieres')

    context = {
        'filieres': filieres,
        'departements': departements
    }
    return render(request, 'administration/liste_filieres.html', context)

def supprimer_filiere(request, pk):
    filiere = get_object_or_404(Filiere, pk=pk)
    if filiere.classes.exists():
        messages.error(request, "Impossible de supprimer : des classes sont liees a cette filiere.")
    else:
        filiere.delete()
        messages.warning(request, "Filiere supprimee.")
    return redirect('administration:liste_filieres')

def modifier_filiere(request, pk):
    filiere = get_object_or_404(Filiere, pk=pk)
    if request.method == 'POST':
        nouveau_nom = request.POST.get('nom').strip()
        nouveau_code = request.POST.get('code').strip().upper()
        nouveau_dept_id = request.POST.get('departement')
        
        # Vérification d'unicité (exclure la filière actuelle de la recherche)
        if Filiere.objects.filter(nom__iexact=nouveau_nom).exclude(pk=pk).exists() or \
           Filiere.objects.filter(code__iexact=nouveau_code).exclude(pk=pk).exists():
            messages.error(request, "Une autre filière possède déjà ce nom ou ce code.")
        else:
            nouveau_dept = get_object_or_404(Departement, id=nouveau_dept_id)
            filiere.nom = nouveau_nom
            filiere.code = nouveau_code
            filiere.departement = nouveau_dept
            filiere.save()
            messages.success(request, f"La filière {filiere.code} a été mise à jour.")
            
    return redirect('administration:liste_filieres')



def liste_classes(request):
    classes = Classe.objects.all().select_related('filiere').order_by('code')
    filieres = Filiere.objects.all().order_by('nom')
    
    if request.method == 'POST':
        nom = request.POST.get('nom').strip()
        code = request.POST.get('code').strip().upper()
        filiere_id = request.POST.get('filiere')
        
        if Classe.objects.filter(code__iexact=code).exists():
            messages.error(request, f"La classe avec le code {code} existe déjà.")
        else:
            filiere = get_object_or_404(Filiere, id=filiere_id)
            Classe.objects.create(nom=nom, code=code, filiere=filiere)
            messages.success(request, f"Classe {code} créée.")
        return redirect('administration:liste_classes')

    return render(request, 'administration/liste_classes.html', {'classes': classes, 'filieres': filieres})

def modifier_classe(request, pk):
    classe = get_object_or_404(Classe, pk=pk)
    if request.method == 'POST':
        classe.nom = request.POST.get('nom').strip()
        classe.code = request.POST.get('code').strip().upper()
        classe.filiere = get_object_or_404(Filiere, id=request.POST.get('filiere'))
        classe.save()
        messages.success(request, "Classe mise à jour.")
    return redirect('administration:liste_classes')
#vague
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Vague, AnneeScolaire

@login_required
def liste_vagues(request):
    maintenant = timezone.now()
    user = request.user
    
    # Sécurité : Seul un Chef de Dept ou le DE peut accéder à cette vue
    if user.role not in ['CHEF_DEPT', 'DE']:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    # --- 1. FILTRAGE DES DONNÉES PAR DÉPARTEMENT ---
    # Si c'est un Chef, il ne voit que son département
    # Si c'est le DE, il voit tout (pour supervision)
    if user.role == 'CHEF_DEPT':
        vagues_query = Vague.objects.filter(departement=user.departement)
    else:
        vagues_query = Vague.objects.all()

    # --- 2. CLÔTURE AUTOMATIQUE (Ciblée) ---
    vagues_query.filter(
        date_fermeture__lte=maintenant, 
        est_cloturee=False
    ).update(est_cloturee=True)

    # Récupération finale pour l'affichage
    vagues = vagues_query.prefetch_related('annees_concernees').order_by('-date_ouverture')
    toutes_les_annees = AnneeScolaire.objects.all().order_by('-libelle')
    
    # Vérification d'une vague active UNIQUEMENT dans ce département
    vague_en_cours = vagues_query.filter(est_cloturee=False).first()

    # --- 3. CRÉATION (Seulement pour le Chef de Dept) ---
    if request.method == 'POST' and user.role == 'CHEF_DEPT':
        libelle = request.POST.get('libelle')
        ouverture_str = request.POST.get('date_ouverture')
        fermeture_str = request.POST.get('date_fermeture')
        annees_ids = request.POST.getlist('annees_concernees')

        ouverture = timezone.make_aware(timezone.datetime.strptime(ouverture_str, '%Y-%m-%dT%H:%M'))
        fermeture = timezone.make_aware(timezone.datetime.strptime(fermeture_str, '%Y-%m-%dT%H:%M'))

        # VALIDATIONS LOCALES AU DÉPARTEMENT
        if vague_en_cours:
            messages.error(request, f"Interdit : Une session est déjà active pour le département {user.departement.nom}.")
            return redirect('administration:liste_vagues')

        if fermeture <= ouverture:
            messages.error(request, "La fermeture doit être après l'ouverture.")
            return redirect('administration:liste_vagues')

        # CRÉATION Rattachée au département
        nouvelle_vague = Vague.objects.create(
            libelle=libelle, 
            date_ouverture=ouverture, 
            date_fermeture=fermeture,
            departement=user.departement, # CRUCIAL
            est_cloturee=False
        )
        nouvelle_vague.annees_concernees.set(annees_ids)
        
        messages.success(request, f"Vague '{libelle}' créée pour votre département.")
        return redirect('administration:liste_vagues')

    return render(request, 'administration/liste_vagues.html', {
        'vagues': vagues, 
        'toutes_les_annees': toutes_les_annees,
        'vague_en_cours': vague_en_cours,
        'maintenant': maintenant
    })
@login_required
def cloturer_vague(request, pk):
    # Sécurité : On récupère la vague seulement si elle appartient au département du chef
    vague = get_object_or_404(Vague, pk=pk, departement=request.user.departement)
    
    if request.method == 'POST':
        vague.est_cloturee = True
        vague.save()
        messages.warning(request, f"La session '{vague.libelle}' est désormais clôturée.")
    
    return redirect('administration:liste_vagues')

@login_required
def modifier_vague(request, pk):
    vague = get_object_or_404(Vague, pk=pk, departement=request.user.departement)
    maintenant = timezone.now()
    
    if request.method == 'POST':
        libelle = request.POST.get('libelle')
        ouverture = request.POST.get('date_ouverture')
        fermeture = request.POST.get('date_fermeture')
        annees_ids = request.POST.getlist('annees_concernees')
        
        # Validation : Existe-t-il UNE AUTRE vague active dans MON département ?
        vague_conflit = Vague.objects.filter(
            departement=request.user.departement,
            est_cloturee=False
        ).exclude(pk=vague.pk).first()

        if vague_conflit:
            messages.error(request, f"Conflit : La session '{vague_conflit.libelle}' est déjà active pour votre département.")
            return redirect('administration:liste_vagues')

        # Mise à jour
        vague.libelle = libelle
        vague.date_ouverture = ouverture
        vague.date_fermeture = fermeture
        vague.save()
        vague.annees_concernees.set(annees_ids)
        
        messages.success(request, f"La vague '{vague.libelle}' a été mise à jour.")
    
    return redirect('administration:liste_vagues')

@login_required
def reouvrir_vague(request, pk):
    vague = get_object_or_404(Vague, pk=pk, departement=request.user.departement)
    maintenant = timezone.now()

    if request.method == 'POST':
        nouvelle_fermeture_str = request.POST.get('nouvelle_date_fermeture')
        nouvelle_fermeture = timezone.make_aware(timezone.datetime.strptime(nouvelle_fermeture_str, '%Y-%m-%dT%H:%M'))

        # Sécurité : Pas d'autre vague active dans le département
        vague_active = Vague.objects.filter(
            departement=request.user.departement,
            est_cloturee=False
        ).exclude(pk=vague.pk).first()

        if vague_active:
            messages.error(request, f"Impossible : '{vague_active.libelle}' est déjà active.")
            return redirect('administration:liste_vagues')

        if nouvelle_fermeture <= maintenant:
            messages.error(request, "La date doit être dans le futur.")
            return redirect('administration:liste_vagues')

        vague.date_fermeture = nouvelle_fermeture
        vague.est_cloturee = False
        vague.save()
        messages.success(request, f"La vague '{vague.libelle}' est de nouveau ouverte.")
    
    return redirect('administration:liste_vagues')

#users
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import get_user_model
from app_administration.models import Departement

User = get_user_model()
# app_administration/views.py
def liste_chefs(request):
    chefs = User.objects.filter(role='CHEF_DEPT').order_by('-date_joined')
    departements = Departement.objects.all()

    if request.method == 'POST':
        # --- ACTION : CRÉATION ---
        if 'creer_chef' in request.POST:
            password = request.POST.get('password')
            if len(password) < 8:
                messages.error(request, "Échec : Le mot de passe doit faire au moins 8 caractères.")
                return redirect('administration:liste_chefs')
            
            user = User.objects.create_user(
                username=request.POST.get('username'),
                email=request.POST.get('email'),
                password=password,
                role='CHEF_DEPT',
                departement_id=request.POST.get('departement')
            )
            messages.success(request, f"Compte de {user.username} créé.")

        # --- ACTION : MODIFIER INFOS (Login, Email, Dept) ---
        elif 'modifier_infos' in request.POST:
            u = get_object_or_404(User, id=request.POST.get('user_id'))
            u.username = request.POST.get('username')
            u.email = request.POST.get('email')
            u.departement_id = request.POST.get('departement')
            u.save()
            messages.success(request, "Informations mises à jour.")

        # --- ACTION : MODIFIER PASSWORD ---
        elif 'modifier_password' in request.POST:
            u = get_object_or_404(User, id=request.POST.get('user_id'))
            nouveau_pass = request.POST.get('nouveau_password')
            if len(nouveau_pass) < 8:
                messages.error(request, "Le nouveau mot de passe est trop court.")
            else:
                u.set_password(nouveau_pass)
                u.save()
                messages.success(request, "Mot de passe modifié.")

        return redirect('administration:liste_chefs')

    return render(request, 'administration/liste_chefs.html', {'chefs': chefs, 'departements': departements})

def supprimer_chef(request, pk):
    chef = get_object_or_404(User, pk=pk)
    nom = chef.username
    chef.delete()
    messages.success(request, f"Le compte de {nom} a été supprimé.")
    return redirect('administration:liste_chefs')

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import Q, Count
from django.utils import timezone
@login_required
def accueil_chef(request):
    maintenant = timezone.now()
    is_de = (request.user.role == 'DE')
    
    # --- LOGIQUE DE SÉLECTION DU DÉPARTEMENT ---
    if is_de:
        dept_id = request.GET.get('dept_id')
        if not dept_id:
            # Si le DE n'a pas choisi, on affiche la liste des départements
            departements = Departement.objects.all().annotate(
                total_etudiants=models.Count('filieres__classes__etudiant')
            )
            return render(request, 'administration/supervision/choix_departement.html', {
                'departements': departements
            })
        # Si un ID est fourni, le DE travaille sur ce département spécifique
        chef_dept = get_object_or_404(Departement, id=dept_id)
    else:
        # Chef de département classique : limité à son propre département
        chef_dept = request.user.departement

    # --- 1. BASE DES DOSSIERS (Filtrée par le département choisi) ---
    dossiers_dept = DossierMemoire.objects.filter(
        etudiant__classe__filiere__departement=chef_dept
    )

    # --- 2. IDENTIFIER LA VAGUE ACTIVE DU DÉPARTEMENT ---
    vague_active = Vague.objects.filter(
        departement=chef_dept,
        date_ouverture__lte=maintenant,
        date_fermeture__gte=maintenant,
        est_cloturee=False
    ).first()

    # --- 3. COMPTEURS DE TRAVAIL ---
    attente_proposition = dossiers_dept.filter(
        Q(soutenance__isnull=True) | Q(soutenance__status='REJETE')
    ).count()

    attente_programmation = dossiers_dept.filter(
        soutenance__status='VALIDE'
    ).count()

    # --- 4. CONSTRUCTION DU CONTEXTE ---
    context = {
        'chef_dept': chef_dept,
        'is_de': is_de,
        'dossiers_count': dossiers_dept.count(),
        'dossiers_hors_vague': dossiers_dept.filter(vague__isnull=True).count(),
        'total_etudiants_dept': Etudiant.objects.filter(
            classe__filiere__departement=chef_dept
        ).count(),
        'vague_active': vague_active,
        'attente_proposition': attente_proposition,
        'attente_programmation': attente_programmation,
        'derniers_etudiants': dossiers_dept.order_by('-id')[:5],
    }

    # Choix du template : le DE garde son interface Superadmin
    template = 'administration/supervision/accueil_miroir.html' if is_de else 'administration/chef/accueil.html'
    return render(request, template, context)
from django.utils import timezone
from .models import AnneeScolaire
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from django.db.models import Q
from django.core.exceptions import FieldError
from django.http import JsonResponse

from .models import Filiere, Classe, Vague, AnneeScolaire
from app_gestion_interne.models import Etudiant, DossierMemoire
@login_required
def liste_themes_etudiants(request):
    is_de = (request.user.role == 'DE')
    
    # --- DÉTERMINATION DU DÉPARTEMENT ---
    if is_de:
        dept_id = request.GET.get('dept_id')
        if not dept_id:
            return redirect('administration:accueil_chef')
        chef_dept = get_object_or_404(Departement, id=dept_id)
    else:
        chef_dept = request.user.departement

    # 1. Filtres GET
    f_id = request.GET.get('filiere')
    c_id = request.GET.get('classe')
    a_id = request.GET.get('annee')
    search = request.GET.get('search')

    # 2. Querybase : Filtrée par le département sélectionné
    dossiers = DossierMemoire.objects.filter(
        etudiant__classe__filiere__departement=chef_dept
    ).exclude(
        Q(theme__isnull=True) | Q(theme="")
    ).select_related('etudiant', 'etudiant__classe', 'etudiant__annee')

    # Application des filtres de recherche
    if f_id: dossiers = dossiers.filter(etudiant__classe__filiere_id=f_id)
    if c_id: dossiers = dossiers.filter(etudiant__classe_id=c_id)
    if a_id: dossiers = dossiers.filter(etudiant__annee_id=a_id)
    if search:
        dossiers = dossiers.filter(
            Q(etudiant__nom__icontains=search) | 
            Q(etudiant__prenom__icontains=search) |
            Q(etudiant__matricule__icontains=search)
        )

    context = {
        'dossiers': dossiers.order_by('-id'),
        'chef_dept': chef_dept,
        'is_de': is_de,
        'filieres': Filiere.objects.filter(departement=chef_dept),
        'annees_scolaires': AnneeScolaire.objects.all().order_by('-libelle'),
        'classes_du_dept': Classe.objects.filter(filiere__departement=chef_dept),
    }

    # Template différencié pour garder la sidebar bleue du DE
    template = 'administration/supervision/themes_supervision.html' if is_de else 'administration/chef/themes_liste.html'
    return render(request, template, context)
from django.http import JsonResponse

def get_classes_by_filiere(request, filiere_id):
    """Retourne les classes d'une filière donnée"""
    classes = Classe.objects.filter(filiere_id=filiere_id).values('id', 'nom')
    return JsonResponse(list(classes), safe=False)

def get_etudiants_par_classe_et_annee(request, classe_id):
    """Retourne les étudiants d'une classe qui n'ont PAS encore de dossier cette année"""
    # On récupère l'année en cours ou on laisse le filtrage global
    etudiants = Etudiant.objects.filter(classe_id=classe_id).values('matricule', 'nom', 'prenom')
    return JsonResponse(list(etudiants), safe=False)


@login_required
def enregistrer_theme_etudiant(request):
    if request.method == 'POST':
        # 1. Récupération des données du formulaire
        dept_id = request.POST.get('dept_id')  # Utilisé par le DE pour rester sur le même département
        dossier_id = request.POST.get('dossier_id')  # Présent si c'est une modification
        matricule = request.POST.get('etudiant_matricule')
        theme = request.POST.get('theme')
        encadreur = request.POST.get('encadreur')
        lieu = request.POST.get('lieu_stage')

        # 2. Récupération de l'étudiant
        etudiant = get_object_or_404(Etudiant, matricule=matricule)

        # 3. Recherche d'un dossier existant pour cet étudiant
        dossier = DossierMemoire.objects.filter(etudiant=etudiant).first()

        if dossier:
            # Si le dossier existe, on autorise la modification si :
            # - Le thème n'était pas encore validé
            # - OU si on vient explicitement d'un formulaire de modification (dossier_id présent)
            if not dossier.is_theme_valide or dossier_id:
                dossier.theme = theme
                dossier.encadreur = encadreur
                dossier.lieu_stage = lieu
                dossier.is_theme_valide = True
                dossier.save()
                messages.success(request, f"Le thème de {etudiant.nom} {etudiant.prenom} a été mis à jour et validé.")
            else:
                # Sécurité pour éviter d'écraser un thème déjà validé par erreur
                messages.error(request, f"Action refusée : {etudiant.nom} possède déjà un thème validé.")
        else:
            # Si aucun dossier n'existe, on le crée de zéro
            DossierMemoire.objects.create(
                etudiant=etudiant,
                theme=theme,
                encadreur=encadreur,
                lieu_stage=lieu,
                is_theme_valide=True
            )
            messages.success(request, f"Dossier créé et thème enregistré pour {etudiant.nom} {etudiant.prenom}.")

        # --- REDIRECTION INTELLIGENTE ---
        # Si c'est le DE, on le renvoie sur la page de supervision du département spécifique
        if request.user.role == 'DE' and dept_id:
            return redirect(f"{reverse('administration:liste_themes_etudiants')}?dept_id={dept_id}")

    # Redirection par défaut (pour les Chefs de Département ou si pas de dept_id)
    return redirect('administration:liste_themes_etudiants')


@login_required
def get_etudiants_par_classe_et_annee(request, classe_id, annee_id):
    # On récupère les étudiants qui appartiennent à cette classe ET à cette année
    etudiants = Etudiant.objects.filter(
        classe_id=classe_id,
        annee_id=annee_id
    ).values('matricule', 'nom', 'prenom')

    # Transformation en liste pour le JsonResponse
    return JsonResponse(list(etudiants), safe=False)
@login_required
def get_classes(request, filiere_id):
    classes = Classe.objects.filter(filiere_id=filiere_id).values('id', 'nom')
    return JsonResponse(list(classes), safe=False)
@login_required
@login_required
def supprimer_dossier_chef(request, pk):
    dossier = get_object_or_404(DossierMemoire, pk=pk)
    # On sauvegarde le dept_id avant la suppression pour la redirection
    dept_id = dossier.etudiant.classe.filiere.departement.id
    
    nom_etudiant = f"{dossier.etudiant.nom} {dossier.etudiant.prenom}"
    dossier.delete()
    messages.success(request, f"Dossier de {nom_etudiant} supprimé.")

    if request.user.role == 'DE':
        return redirect(f"{reverse('administration:liste_themes_etudiants')}?dept_id={dept_id}")
    return redirect('administration:liste_themes_etudiants')


from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Vague, AnneeScolaire

from django.db.models import Q

from django.utils.dateparse import parse_datetime
from django.utils import timezone


@login_required
def chef_liste_vagues(request):
    maintenant = timezone.now()
    is_de = (request.user.role == 'DE')

    # --- LOGIQUE DE DÉTERMINATION DU DÉPARTEMENT ---
    if is_de:
        # Le DE choisit le département via l'URL (?dept_id=...)
        dept_id = request.GET.get('dept_id') or request.POST.get('dept_id')
        if not dept_id:
            return redirect('administration:accueil_chef')
        chef_dept = get_object_or_404(Departement, id=dept_id)
    else:
        # Le Chef est bloqué sur son département à lui
        chef_dept = request.user.departement

    # --- RÉCUPÉRATION DES DONNÉES ---
    vagues_query = Vague.objects.filter(departement=chef_dept)

    # Filtres de recherche
    search_libelle = request.GET.get('search_vague')
    search_annee = request.GET.get('search_annee_crea')
    if search_libelle:
        vagues_query = vagues_query.filter(libelle__icontains=search_libelle)
    if search_annee:
        vagues_query = vagues_query.filter(date_creation__year=search_annee)

    vagues = vagues_query.prefetch_related('annees_concernees').order_by('-date_creation')

    # Données pour les modales et l'affichage
    toutes_les_annees = AnneeScolaire.objects.all().order_by('-libelle')
    annees_existantes = Vague.objects.filter(departement=chef_dept).dates('date_creation', 'year', order='DESC')
    vague_en_cours = Vague.objects.filter(departement=chef_dept, est_cloturee=False).first()

    # --- GESTION DE LA CRÉATION (POST) ---
    if request.method == 'POST':
        if vague_en_cours:
            messages.error(request, "Impossible : une session est déjà active pour ce département.")
        else:
            libelle = request.POST.get('libelle')
            ouverture_str = request.POST.get('date_ouverture')
            fermeture_str = request.POST.get('date_fermeture')

            # Validation simple des dates
            if ouverture_str and fermeture_str and fermeture_str <= ouverture_str:
                messages.error(request, "Erreur : La date de fermeture doit être après l'ouverture.")
            else:
                v = Vague.objects.create(
                    libelle=libelle,
                    departement=chef_dept,
                    date_ouverture=ouverture_str,
                    date_fermeture=fermeture_str,
                    est_cloturee=False
                )
                v.annees_concernees.set(request.POST.getlist('annees_concernees'))
                messages.success(request, f"Session créée avec succès pour {chef_dept.nom}.")

                # Redirection intelligente : on garde le dept_id si on est DE
                if is_de:
                    return redirect(f"{reverse('administration:chef_liste_vagues')}?dept_id={chef_dept.id}")
                return redirect('administration:chef_liste_vagues')

    context = {
        'vagues': vagues,
        'chef_dept': chef_dept,
        'is_de': is_de,
        'toutes_les_annees': toutes_les_annees,
        'annees_existantes': annees_existantes,
        'vague_en_cours': vague_en_cours
    }

    template = 'administration/supervision/vagues_supervision.html' if is_de else 'administration/chef/vagues_liste.html'
    return render(request, template, context)
# --- FONCTIONS D'ACTIONS (Appelées par les modales) ---
@login_required
def modifier_vague(request, pk):
    # 1. On récupère la vague sans filtrer immédiatement par request.user.departement
    vague = get_object_or_404(Vague, pk=pk)
    is_de = (request.user.role == 'DE')

    # 2. Sécurité : Vérifier si l'utilisateur a le droit (Soit DE, soit le Chef du bon département)
    if not is_de and request.user.departement != vague.departement:
        messages.error(request, "Accès non autorisé.")
        return redirect('administration:accueil_chef')

    if request.method == 'POST':
        vague.libelle = request.POST.get('libelle')
        vague.date_ouverture = request.POST.get('date_ouverture')
        vague.date_fermeture = request.POST.get('date_fermeture')
        vague.save()

        # Mise à jour ManyToMany
        ids_selectionnes = request.POST.getlist('annees_concernees')
        vague.annees_concernees.set(ids_selectionnes)

        messages.success(request, f"Session '{vague.libelle}' mise à jour.")

    # 3. Redirection intelligente selon le rôle
    if is_de:
        return redirect(f"{reverse('administration:chef_liste_vagues')}?dept_id={vague.departement.id}")
    return redirect('administration:chef_liste_vagues')


@login_required
def cloturer_vague(request, pk):
    vague = get_object_or_404(Vague, pk=pk)
    is_de = (request.user.role == 'DE')

    if not is_de and request.user.departement != vague.departement:
        messages.error(request, "Accès non autorisé.")
        return redirect('administration:accueil_chef')

    vague.est_cloturee = True
    vague.save()
    messages.warning(request, f"Session '{vague.libelle}' archivée.")

    if is_de:
        return redirect(f"{reverse('administration:chef_liste_vagues')}?dept_id={vague.departement.id}")
    return redirect('administration:chef_liste_vagues')


@login_required
def reouvrir_vague(request, pk):
    vague = get_object_or_404(Vague, pk=pk)
    is_de = (request.user.role == 'DE')
    dept_cible = vague.departement  # On utilise le département lié à la vague

    if not is_de and request.user.departement != dept_cible:
        messages.error(request, "Accès non autorisé.")
        return redirect('administration:accueil_chef')

    if request.method == 'POST':
        # Vérifier si une AUTRE vague est déjà active dans CE département
        vague_active = Vague.objects.filter(
            departement=dept_cible,
            est_cloturee=False
        ).exclude(pk=vague.pk).exists()

        if vague_active:
            messages.error(request, "Action refusée : Une autre session est déjà active pour ce département.")
        else:
            nouvelle_date_str = request.POST.get('nouvelle_date_fermeture')
            if nouvelle_date_str:
                nouvelle_date_naive = parse_datetime(nouvelle_date_str)
                if nouvelle_date_naive:
                    nouvelle_date = timezone.make_aware(nouvelle_date_naive)

                    if nouvelle_date <= vague.date_fermeture:
                        messages.error(request, "La nouvelle date doit être après l'ancienne clôture.")
                    else:
                        vague.date_fermeture = nouvelle_date
                        vague.est_cloturee = False
                        vague.save()
                        messages.success(request, f"La session '{vague.libelle}' est de nouveau ouverte.")

    if is_de:
        return redirect(f"{reverse('administration:chef_liste_vagues')}?dept_id={dept_cible.id}")
    return redirect('administration:chef_liste_vagues')
from django.db.models import Q
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required


@login_required
def chef_detail_vague(request, pk):
    """
    Détail d'une session : affiche les dossiers inscrits et permet l'inscription
    de nouveaux dossiers éligibles (pré-dépôt validé et promotion correspondante).
    """
    # 1. Récupération de la vague (sans filtre département pour permettre l'accès au DE)
    vague = get_object_or_404(Vague, pk=pk)
    is_de = (request.user.role == 'DE')
    maintenant = timezone.now()

    # 2. LOGIQUE DE DÉTERMINATION DU DÉPARTEMENT & SÉCURITÉ
    if is_de:
        # Le DE "emprunte" le département de la vague qu'il consulte
        chef_dept = vague.departement
    else:
        # Le Chef est strictement limité à son département
        if vague.departement != request.user.departement:
            messages.error(request, "Accès non autorisé à ce département.")
            return redirect('administration:accueil_chef')
        chef_dept = request.user.departement

    # 3. BASE DES DOSSIERS ÉLIGIBLES
    # Doivent avoir validé le pré-dépôt et appartenir au département
    dossiers_base = DossierMemoire.objects.filter(
        etudiant__classe__filiere__departement=chef_dept,
        is_pre_depot_fait=True
    )

    # Filtrage strict par les promotions (Années Scolaires) autorisées pour cette vague
    annees_valides_ids = vague.annees_concernees.values_list('id', flat=True)
    dossiers_base = dossiers_base.filter(etudiant__annee_id__in=annees_valides_ids)

    # 4. LOGIQUE D'AFFICHAGE & ÉTAT DE LA VAGUE
    # Une vague est modifiable si elle est dans son créneau de dates ET non clôturée manuellement
    peut_modifier = vague.date_ouverture <= maintenant <= vague.date_fermeture and not vague.est_cloturee
    peut_reouvrir = vague.date_ouverture <= maintenant <= vague.date_fermeture and vague.est_cloturee

    if peut_modifier:
        # On exclut les étudiants déjà inscrits dans une AUTRE vague qui est encore OUVERTE
        # pour éviter les doublons d'inscriptions actives
        occupes_ailleurs = DossierMemoire.objects.filter(
            vague__departement=chef_dept,
            vague__est_cloturee=False,
            is_soutenu=False
        ).exclude(vague=vague).values_list('etudiant_id', flat=True)

        # On affiche : Les dossiers déjà dans cette vague + Ceux non soutenus libres
        dossiers = dossiers_base.filter(
            Q(vague=vague) | Q(is_soutenu=False)
        ).exclude(etudiant_id__in=occupes_ailleurs)
    else:
        # Si la session est close ou expirée : on affiche uniquement les inscrits définitifs
        dossiers = dossiers_base.filter(vague=vague)

    # 5. FILTRES DE RECHERCHE (Filière, Classe, Année)
    f_id = request.GET.get('filiere')
    c_id = request.GET.get('classe')
    a_id = request.GET.get('annee')

    if f_id and f_id.isdigit():
        dossiers = dossiers.filter(etudiant__classe__filiere_id=f_id)
    if c_id and c_id.isdigit():
        dossiers = dossiers.filter(etudiant__classe_id=c_id)
    if a_id and a_id.isdigit():
        dossiers = dossiers.filter(etudiant__annee_id=a_id)

    # Optimisation des requêtes pour éviter le N+1
    dossiers = dossiers.select_related(
        'etudiant',
        'etudiant__classe',
        'etudiant__annee',
        'etudiant__classe__filiere'
    ).distinct().order_by('etudiant__nom')

    # 6. CONTEXTE POUR LE TEMPLATE
    context = {
        'vague': vague,
        'dossiers': dossiers,
        'chef_dept': chef_dept,
        'is_de': is_de,
        'peut_modifier': peut_modifier,
        'peut_reouvrir': peut_reouvrir,
        'filieres': Filiere.objects.filter(departement=chef_dept),
        'classes_du_dept': Classe.objects.filter(filiere__departement=chef_dept),
        'maintenant': maintenant,
    }

    # 7. ROUTAGE DYNAMIQUE DU TEMPLATE
    # Chaque rôle garde son environnement (Sidebar, Navbar, Styles)
    if is_de:
        return render(request, 'administration/supervision/vague_detail_supervision.html', context)
    return render(request, 'administration/chef/vague_detail.html', context)


@login_required
def chef_toggle_cloture(request, pk):
    """
    Inverse l'état de clôture manuelle de la session.
    Adapté pour le DE (supervision multi-départements) et le Chef.
    """
    # 1. Récupération sans filtrer par request.user.departement pour le DE
    vague = get_object_or_404(Vague, pk=pk)
    is_de = (request.user.role == 'DE')
    maintenant = timezone.now()

    # 2. Sécurité : Si pas DE, vérification du département
    if not is_de and vague.departement != request.user.departement:
        messages.error(request, "Accès non autorisé.")
        return redirect('administration:accueil_chef')

    chef_dept = vague.departement

    # 3. Vérification du délai global (Optionnel : le DE pourrait avoir le droit de passer outre ?)
    # Ici, on garde la règle académique : si la date de fermeture est passée, on ne peut plus ouvrir.
    if maintenant > vague.date_fermeture:
        messages.error(request, "Impossible : le délai de fermeture académique est dépassé.")
        return self_redirect_vague(vague.id, chef_dept.id, is_de)

    # 4. Logique de basculement (Toggle)
    if vague.est_cloturee:
        # TENTATIVE DE RÉOUVERTURE
        # Vérification s'il existe une AUTRE vague déjà ouverte dans CE département
        autre_vague_ouverte = Vague.objects.filter(
            departement=chef_dept,
            est_cloturee=False
        ).exclude(pk=vague.pk).exists()

        if autre_vague_ouverte:
            messages.error(
                request,
                f"Action refusée : Une autre session est déjà active pour le département {chef_dept.nom}."
            )
            return self_redirect_vague(vague.id, chef_dept.id, is_de)

        vague.est_cloturee = False
        msg = "réouverte"
        messages.success(request, f"La session a été {msg} avec succès.")
    else:
        # CLÔTURE MANUELLE
        vague.est_cloturee = True
        msg = "clôturée"
        messages.warning(request, f"La session a été {msg} (archivée).")

    vague.save()
    return self_redirect_vague(vague.id, chef_dept.id, is_de)


# Utilitaire de redirection pour cohérence
def self_redirect_vague(vague_id, dept_id, is_de):
    url = reverse('administration:chef_detail_vague', kwargs={'pk': vague_id})
    if is_de:
        return redirect(f"{url}?dept_id={dept_id}")
    return redirect(url)
@login_required
def chef_inscrire_vague(request, pk_dossier, pk_vague):
    """Inscrit un étudiant à la vague (Adapté Chef et DE)"""
    vague = get_object_or_404(Vague, pk=pk_vague)
    is_de = (request.user.role == 'DE')

    # --- SÉCURITÉ ---
    if not is_de and vague.departement != request.user.departement:
        messages.error(request, "Accès non autorisé.")
        return redirect('administration:accueil_chef')

    dossier = get_object_or_404(DossierMemoire, pk=pk_dossier)
    maintenant = timezone.now()

    # Vérification si la vague est encore ouverte (Dates + État)
    if not vague.est_cloturee and vague.date_ouverture <= maintenant <= vague.date_fermeture:
        dossier.vague = vague
        dossier.save()
        messages.success(request, f"{dossier.etudiant.nom} a été inscrit avec succès.")
    else:
        messages.error(request, "Impossible : cette session est clôturée ou le délai est expiré.")

    # Redirection vers le détail (la vue chef_detail_vague s'occupe de choisir le bon template)
    url = reverse('administration:chef_detail_vague', kwargs={'pk': pk_vague})
    # Si c'est le DE, on ajoute l'ID du département dans l'URL pour garder le menu actif
    if is_de:
        return redirect(f"{url}?dept_id={vague.departement.id}")
    return redirect(url)


@login_required
def chef_annuler_inscription(request, pk_dossier, pk_vague):
    """Retire un étudiant de la vague (Adapté Chef et DE)"""
    vague = get_object_or_404(Vague, pk=pk_vague)
    is_de = (request.user.role == 'DE')

    # --- SÉCURITÉ ---
    if not is_de and vague.departement != request.user.departement:
        messages.error(request, "Accès non autorisé.")
        return redirect('administration:accueil_chef')

    dossier = get_object_or_404(DossierMemoire, pk=pk_dossier)

    # On ne peut annuler que si la vague n'est pas encore clôturée manuellement
    if not vague.est_cloturee:
        dossier.vague = None
        dossier.save()
        messages.info(request, f"L'inscription de {dossier.etudiant.nom} a été annulée.")
    else:
        messages.error(request, "Modification impossible sur une session archivée.")

    url = reverse('administration:chef_detail_vague', kwargs={'pk': pk_vague})
    if is_de:
        return redirect(f"{url}?dept_id={vague.departement.id}")
    return redirect(url)
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Alignment, PatternFill
from django.http import HttpResponse # <--- C'est cette ligne qui manque

from django.db.models import Q


@login_required
def exporter_rapport_complet(request):
    is_de = (request.user.role == 'DE')
    type_donnees = request.GET.get('type')
    format_type = request.GET.get('format', 'excel')
    filiere_id = request.GET.get('filiere')
    vague_id = request.GET.get('vague')
    annee_id = request.GET.get('annee')
    dept_id = request.GET.get('dept_id')  # Nouveau : pour le DE

    # --- BASE DE RECHERCHE ---
    query = DossierMemoire.objects.all()

    # --- FILTRAGE PAR DÉPARTEMENT (Sécurité DE vs Chef) ---
    if is_de:
        if dept_id:
            query = query.filter(etudiant__classe__filiere__departement_id=dept_id)
    else:
        query = query.filter(etudiant__classe__filiere__departement=request.user.departement)

    # --- LOGIQUE FILTRAGE ET HEADERS ---
    if type_donnees == 'themes':
        filename = "Liste_des_Themes"
        headers = ['MATRICULE', 'ÉTUDIANT', 'CLASSE', 'THÈME COMPLET', 'ENCADREUR']
        if annee_id:
            query = query.filter(etudiant__annee_id=annee_id)
    else:
        filename = "Planning_Soutenances"
        headers = ['DATE', 'ÉTUDIANT', 'CLASSE', 'THÈME COMPLET', 'HEURE', 'SALLE']
        if vague_id:
            query = query.filter(vague_id=vague_id)
        query = query.filter(soutenance__status='PROGRAMME')

    if filiere_id:
        query = query.filter(etudiant__classe__filiere_id=filiere_id)

    query = query.select_related('etudiant', 'etudiant__classe', 'soutenance').distinct()
    dossiers = query.order_by('soutenance__date',
                              'soutenance__heure') if type_donnees == 'planning' else query.order_by('etudiant__nom')

    # --- OPTION A : EXCEL ---
    if format_type == 'excel':
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(headers)

        header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        for d in dossiers:
            if type_donnees == 'planning':
                s = d.soutenance
                ws.append([s.date.strftime('%d/%m/%y'), f"{d.etudiant.nom} {d.etudiant.prenom}", d.etudiant.classe.nom,
                           d.theme, s.heure.strftime('%H:%M'), s.salle])
            else:
                matricule_4 = str(d.etudiant.matricule)[:4] if d.etudiant.matricule else ""
                ws.append(
                    [matricule_4, f"{d.etudiant.nom} {d.etudiant.prenom}", d.etudiant.classe.nom, d.theme, d.encadreur])

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
        wb.save(response)
        return response

    # --- OPTION B : WORD (Format Officiel Niger) ---
    else:
        doc = Document()
        # En-tête ESCEP-NIGER
        section = doc.sections[0]
        header_table = doc.add_table(rows=1, cols=2)
        header_table.width = Inches(6.5)

        left_p = header_table.rows[0].cells[0].paragraphs[0]
        left_p.add_run("RÉPUBLIQUE DU NIGER\n").bold = True
        left_p.add_run("Fraternité - Travail - Progrès\n-----------\nESCEP-NIGER\n").bold = True
        left_p.add_run(
            f"DEPT : {dossiers.first().etudiant.classe.filiere.departement.nom if dossiers.exists() else ''}")

        right_p = header_table.rows[0].cells[1].paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.add_run(f"Niamey, le {timezone.now().strftime('%d/%m/%Y')}")

        doc.add_paragraph("\n")
        titre = doc.add_heading(filename.replace('_', ' ').upper(), level=1)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text

        for d in dossiers:
            row = table.add_row().cells
            if type_donnees == 'planning':
                s = d.soutenance
                row[0].text = s.date.strftime('%d/%m/%y')
                row[1].text = f"{d.etudiant.nom} {d.etudiant.prenom}"
                row[2].text = d.etudiant.classe.nom
                row[3].text = d.theme or ""
                row[4].text = s.heure.strftime('%H:%M')
                row[5].text = s.salle or ""
            else:
                row[0].text = str(d.etudiant.matricule)[:4]
                row[1].text = f"{d.etudiant.nom} {d.etudiant.prenom}"
                row[2].text = d.etudiant.classe.nom
                row[3].text = d.theme or ""
                row[4].text = d.encadreur or ""

        doc.add_paragraph(f"\n\nSigné le {timezone.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph(
            "Le Directeur des Études" if is_de else "Le Chef de Département").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        doc.save(response)
        return response


@login_required
def page_rapport_chef(request):
    is_de = (request.user.role == 'DE')

    if is_de:
        dept_id = request.GET.get('dept_id')
        if not dept_id: return redirect('administration:accueil_chef')
        dept = get_object_or_404(Departement, id=dept_id)
    else:
        dept = request.user.departement

    vague_id = request.GET.get('vague')
    filiere_id = request.GET.get('filiere')

    # On récupère les vagues du département, triées par date de création (la plus récente d'abord)
    # prefetch_related est nécessaire car annees_concernees est un ManyToMany
    vagues = Vague.objects.filter(departement=dept).prefetch_related('annees_concernees').order_by('-date_creation')

    dossiers_preview = None
    if vague_id:
        query = DossierMemoire.objects.filter(
            vague_id=vague_id,
            etudiant__classe__filiere__departement=dept,
            soutenance__status='PROGRAMME'
        ).select_related('etudiant', 'etudiant__classe', 'soutenance')

        if filiere_id:
            query = query.filter(etudiant__classe__filiere_id=filiere_id)
        dossiers_preview = query.order_by('soutenance__date', 'soutenance__heure')

    context = {
        'dept': dept,
        'filieres': Filiere.objects.filter(departement=dept),
        'vagues': vagues,
        'dossiers': dossiers_preview,
        'vague_selectionnee': vague_id,
        'filiere_selectionnee': filiere_id,
        'is_de': is_de,
        'toutes_les_annees': AnneeScolaire.objects.all().order_by('-libelle'),
    }
    return render(request, 'administration/supervision/rapports_supervision.html', context)

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q
from app_gestion_interne.models import Jury
@login_required
def liste_jurys_chef(request):
    is_de = (request.user.role == 'DE')
    
    if is_de:
        dept_id = request.GET.get('dept_id')
        if not dept_id: return redirect('administration:accueil_chef')
        chef_dept = get_object_or_404(Departement, id=dept_id)
    else:
        chef_dept = request.user.departement

    # Filtres
    search = request.GET.get('search')
    statut = request.GET.get('statut')

    jurys = Jury.objects.filter(departement=chef_dept).order_by('nom')

    if search:
        jurys = jurys.filter(Q(nom__icontains=search) | Q(prenom__icontains=search))
    if statut:
        jurys = jurys.filter(statut=statut)

    context = {
        'jurys': jurys,
        'chef_dept': chef_dept,
        'is_de': is_de,
    }
    template = 'administration/supervision/jurys_supervision.html' if is_de else 'administration/chef/liste_jurys.html'
    return render(request, template, context)


@login_required
def enregistrer_jury_chef(request):
    """Gère la création et la modification avec support DE."""
    if request.method == 'POST':
        is_de = (request.user.role == 'DE')
        jury_id = request.POST.get('jury_id')

        # Récupération du département cible (via POST pour le DE, via User pour le Chef)
        if is_de:
            dept_id = request.POST.get('dept_id')
            chef_dept = get_object_or_404(Departement, id=dept_id)
        else:
            chef_dept = request.user.departement

        # Nettoyage des données
        nom = request.POST.get('nom').upper().strip()
        prenom = request.POST.get('prenom').strip()
        nouvelle_spec = request.POST.get('specialite').strip()
        statut = request.POST.get('statut')
        telephone = request.POST.get('telephone').strip()

        # --- CAS MODIFICATION ---
        if jury_id:
            # Sécurité : on vérifie que le jury appartient bien au département cible
            jury = get_object_or_404(Jury, id=jury_id, departement=chef_dept)
            jury.nom = nom
            jury.prenom = prenom
            jury.specialite = nouvelle_spec
            jury.statut = statut
            jury.telephone = telephone
            jury.save()
            messages.success(request, f"Le profil de {prenom} {nom} a été mis à jour.")

        # --- CAS NOUVEL AJOUT ---
        else:
            jury_existant = Jury.objects.filter(
                departement=chef_dept,
                nom__iexact=nom,
                prenom__iexact=prenom
            ).first()

            if jury_existant:
                specs_actuelles = [s.strip().lower() for s in jury_existant.specialite.split(',')]
                if nouvelle_spec.lower() not in specs_actuelles:
                    jury_existant.specialite += f", {nouvelle_spec}"
                    jury_existant.save()
                    messages.info(request, f"Spécialité ajoutée au profil de {prenom} {nom}.")
                else:
                    messages.warning(request, f"{prenom} {nom} existe déjà avec cette spécialité.")
            else:
                Jury.objects.create(
                    departement=chef_dept,
                    nom=nom,
                    prenom=prenom,
                    specialite=nouvelle_spec,
                    statut=statut,
                    telephone=telephone
                )
                messages.success(request, "Nouveau membre du jury ajouté.")

        # Redirection intelligente
        url = reverse('administration:liste_jurys_chef')
        return redirect(f"{url}?dept_id={chef_dept.id}" if is_de else url)

    return redirect('administration:liste_jurys_chef')

@login_required
def supprimer_jury_chef(request, pk):
    """Supprime un jury avec sécurité par rôle."""
    jury = get_object_or_404(Jury, pk=pk)
    is_de = (request.user.role == 'DE')
    dept_id = jury.departement.id

    # Sécurité : Le Chef ne peut supprimer que chez lui
    if not is_de and jury.departement != request.user.departement:
        messages.error(request, "Action non autorisée.")
        return redirect('administration:accueil_chef')

    nom_complet = f"{jury.prenom} {jury.nom}"
    jury.delete()
    messages.success(request, f"Le membre {nom_complet} a été supprimé.")

    url = reverse('administration:liste_jurys_chef')
    return redirect(f"{url}?dept_id={dept_id}" if is_de else url)


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from app_gestion_interne.models import DossierMemoire, Soutenance, Jury
from app_administration.models import Filiere, Vague

# --- ESPACE CHEF DE DEPARTEMENT ---
@login_required
def planning_liste_vagues(request):
    is_de = (request.user.role == 'DE')
    
    if is_de:
        dept_id = request.GET.get('dept_id')
        if not dept_id:
            return redirect('administration:accueil_chef')
        chef_dept = get_object_or_404(Departement, id=dept_id)
    else:
        chef_dept = request.user.departement
    
    vagues = Vague.objects.filter(departement=chef_dept).order_by('-date_creation')
    
    context = {
        'vagues': vagues,
        'chef_dept': chef_dept,
        'is_de': is_de
    }
    # Utilisation du template de supervision pour le DE
    template = 'administration/supervision/planning_supervision_vagues.html' if is_de else 'administration/chef/planning_vagues_liste.html'
    return render(request, template, context)
from django.utils import timezone
from datetime import datetime

from django.db.models import Q
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required

@login_required
def planning_vague_detail(request, pk):
    is_de = (request.user.role == 'DE')
    # On récupère la vague. Si c'est le DE, on ne vérifie pas l'appartenance au département dans le filtre de base
    vague = get_object_or_404(Vague, pk=pk)
    
    # Sécurité : Si c'est un Chef, il ne doit voir que son département
    if not is_de and vague.departement != request.user.departement:
        messages.error(request, "Accès refusé.")
        return redirect('administration:accueil_chef')

    chef_dept = vague.departement
    annees_valides_ids = vague.annees_concernees.values_list('id', flat=True)

    dossiers = DossierMemoire.objects.filter(
        vague=vague,
        etudiant__classe__filiere__departement=chef_dept,
        etudiant__annee_id__in=annees_valides_ids
    ).select_related('etudiant', 'etudiant__classe', 'soutenance')

    context = {
        'vague': vague,
        'dossiers': dossiers.order_by('soutenance__date', 'soutenance__heure'),
        'chef_dept': chef_dept,
        'jurys': Jury.objects.filter(departement=chef_dept),
        'is_de': is_de,
    }
    
    template = 'administration/supervision/planning_supervision_detail.html' if is_de else 'administration/chef/planning_vague_detail.html'
    return render(request, template, context)
from django.urls import reverse


@login_required
def proposer_soutenance(request):
    if request.method == "POST":
        d_id = request.POST.get('dossier_id')
        dossier = get_object_or_404(DossierMemoire, id=d_id)

        vague_id = dossier.vague.id
        dept_id = dossier.etudiant.classe.filiere.departement.id
        is_de = (request.user.role == 'DE')

        # --- SÉCURITÉ ---
        if not is_de and dossier.etudiant.classe.filiere.departement != request.user.departement:
            messages.error(request, "Accès non autorisé.")
            return redirect('administration:accueil_chef')

        if dossier.is_soutenu or dossier.vague.est_cloturee:
            messages.error(request, "Modification impossible (Session close ou déjà soutenu).")
            return self_redirect_with_dept(vague_id, dept_id, is_de)

        # --- TRAITEMENT DU JURY ---
        pres = request.POST.get('president')
        rapp = request.POST.get('rapporteur')
        exam = request.POST.get('examinateur')

        membres = [m for m in [pres, rapp, exam] if m]
        if len(set(membres)) < 2:
            messages.error(request, "Sélectionnez au moins deux membres différents.")
            return self_redirect_with_dept(vague_id, dept_id, is_de)

        soutenance, _ = Soutenance.objects.get_or_create(dossier=dossier)
        soutenance.president_id = pres if pres else None
        soutenance.rapporteur_id = rapp if rapp else None
        soutenance.examinateur_id = exam if exam else None

        # LOGIQUE DE STATUT : Le DE valide directement
        if is_de:
            soutenance.status = 'VALIDE'
            messages.success(request, f"Jury validé directement par la Direction pour {dossier.etudiant.nom}.")
        else:
            soutenance.status = 'PROPOSE'
            messages.success(request, f"Proposition de jury envoyée pour {dossier.etudiant.nom}.")

        soutenance.save()
        return self_redirect_with_dept(vague_id, dept_id, is_de)

    return redirect('administration:accueil_chef')

# Petite fonction utilitaire pour gérer la redirection propre
def self_redirect_with_dept(vague_id, dept_id, is_de):
    url = reverse('administration:planning_vague_detail', kwargs={'pk': vague_id})
    if is_de:
        return redirect(f"{url}?dept_id={dept_id}")
    return redirect(url)


@login_required
def enregistrer_planning(request):
    if request.method == "POST":
        s_id = request.POST.get('soutenance_id')
        soutenance = get_object_or_404(Soutenance, id=s_id)

        vague_id = soutenance.dossier.vague.id
        dept_id = soutenance.dossier.etudiant.classe.filiere.departement.id
        is_de = (request.user.role == 'DE')

        # SÉCURITÉ : Statut requis (Le DE peut programmer même si c'est seulement 'VALIDE')
        autorise = ['VALIDE', 'PROGRAMME']
        if not is_de and soutenance.status not in autorise:
            messages.error(request, "Le jury doit être validé avant la programmation.")
            return self_redirect_with_dept(vague_id, dept_id, is_de)

        # Récupération des données
        date_val = request.POST.get('date')
        heure_val = request.POST.get('heure')
        salle_val = request.POST.get('salle')

        if not date_val or not heure_val or not salle_val:
            messages.error(request, "Tous les champs sont obligatoires.")
            return self_redirect_with_dept(vague_id, dept_id, is_de)

        soutenance.date, soutenance.heure, soutenance.salle = date_val, heure_val, salle_val

        # Le passage en PROGRAMME est automatique si les infos sont remplies
        soutenance.status = 'PROGRAMME'
        soutenance.save()

        messages.success(request, f"Planning confirmé pour {soutenance.dossier.etudiant.nom}.")
        return self_redirect_with_dept(vague_id, dept_id, is_de)

    return redirect('administration:accueil_chef')
from django.views.decorators.http import require_POST
@login_required
@require_POST
def tout_cocher_soutenu(request, vague_id):
    # On récupère la vague pour identifier le département concerné
    vague = get_object_or_404(Vague, id=vague_id)
    chef_dept = vague.departement
    is_de = (request.user.role == 'DE')

    # Sécurité : Si ce n'est pas le DE, on vérifie que le chef appartient au bon département
    if not is_de and chef_dept != request.user.departement:
        messages.error(request, "Accès refusé : ce département ne vous est pas assigné.")
        return redirect('administration:accueil_chef')

    # Mise à jour massive : Uniquement les dossiers de cette vague + ce département + déjà programmés
    dossiers_mis_a_jour = DossierMemoire.objects.filter(
        vague_id=vague_id,
        etudiant__classe__filiere__departement=chef_dept,
        soutenance__status='PROGRAMME'
    ).update(is_soutenu=True)

    if dossiers_mis_a_jour > 0:
        messages.success(request, f"Succès : {dossiers_mis_a_jour} étudiant(s) marqué(s) comme soutenus.")
    else:
        messages.warning(request, "Aucun dossier éligible (statut PROGRAMMÉ) n'a été trouvé.")

    # Redirection intelligente via la fonction utilitaire
    return self_redirect_with_dept(vague_id, chef_dept.id, is_de)
@login_required
def marquer_soutenu_manuel(request, pk_dossier):
    # On récupère le dossier sans filtrer par le département du user pour laisser le DE agir
    dossier = get_object_or_404(DossierMemoire, id=pk_dossier)
    
    # Sécurité : Seul le DE ou le Chef du département concerné peut agir
    if request.user.role != 'DE' and dossier.etudiant.classe.filiere.departement != request.user.departement:
        messages.error(request, "Droit insuffisant.")
        return redirect('administration:accueil_chef')

    if hasattr(dossier, 'soutenance') and dossier.soutenance.status in ['PROGRAMME', 'VALIDE']:
        dossier.is_soutenu = not dossier.is_soutenu
        dossier.save()
        
    return redirect(request.META.get('HTTP_REFERER'))


@login_required
def supprimer_proposition_soutenance(request, pk_soutenance):
    """Adapté pour que le DE puisse aussi supprimer dans n'importe quel département"""
    soutenance = get_object_or_404(Soutenance, id=pk_soutenance)
    is_de = (request.user.role == 'DE')
    dept = soutenance.dossier.etudiant.classe.filiere.departement

    # Sécurité
    if not is_de and dept != request.user.departement:
        messages.error(request, "Droit insuffisant.")
        return redirect('administration:accueil_chef')

    vague_id = soutenance.dossier.vague.id

    if soutenance.dossier.is_soutenu:
        messages.error(request, "Impossible : l'étudiant a déjà soutenu.")
    else:
        soutenance.delete()
        messages.success(request, "Programmation supprimée.")

    return self_redirect_with_dept(vague_id, dept.id, is_de)


# Utilitaire de redirection (à placer dans votre views.py)
def self_redirect_with_dept(vague_id, dept_id, is_de):
    url = reverse('administration:planning_vague_detail', kwargs={'pk': vague_id})
    if is_de:
        return redirect(f"{url}?dept_id={dept_id}")
    return redirect(url)
# --- ESPACE DIRECTION ---
@login_required
def validation_jury_liste(request):
    # Récupérer les paramètres de filtrage
    annee_id = request.GET.get('annee')
    filiere_id = request.GET.get('filiere')
    classe_id = request.GET.get('classe')

    # Base de la requête : uniquement les jurys proposés
    propositions = Soutenance.objects.filter(status='PROPOSE').select_related(
        'dossier__etudiant__classe__filiere__departement', 
        'dossier__vague',
        'president', 'rapporteur'
    )

    # Application des filtres
    if annee_id and annee_id.isdigit():
        propositions = propositions.filter(dossier__vague__annees_concernees__id=annee_id)
    if filiere_id:
        propositions = propositions.filter(dossier__etudiant__classe__filiere_id=filiere_id)
    if classe_id:
        propositions = propositions.filter(dossier__etudiant__classe_id=classe_id)

    context = {
        'propositions': propositions,
        'annees': AnneeScolaire.objects.all().order_by('-libelle'),        'filieres': Filiere.objects.all(),
        # On passe les classes uniquement si une filière est choisie (pour le script JS)
        'classes': Classe.objects.filter(filiere_id=filiere_id) if filiere_id else None,
    }
    return render(request, 'administration/validation_jury_liste.html', context)

@login_required
def traiter_validation_jury(request, pk):
    """Action de valider ou rejeter une proposition"""
    soutenance = get_object_or_404(Soutenance, pk=pk)
    
    if request.method == "POST":
        action = request.POST.get('action') # 'valider' ou 'rejeter'
        motif = request.POST.get('motif_rejet')

        if action == 'valider':
            soutenance.status = 'VALIDE'
            soutenance.motif_rejet = ""
            messages.success(request, f"Jury validé pour {soutenance.dossier.etudiant.nom}.")
        
        elif action == 'rejeter':
            if not motif:
                messages.error(request, "Vous devez fournir un motif pour le rejet.")
                return redirect('administration:validation_jury_liste')
            
            soutenance.status = 'REJETE'
            soutenance.motif_rejet = motif
            messages.warning(request, f"Proposition rejetée pour {soutenance.dossier.etudiant.nom}.")
        
        soutenance.save()
        
    return redirect('administration:validation_jury_liste')


import pandas as pd
from django.shortcuts import render, redirect
from django.contrib import messages

def liste_etudiants_de(request):
    # --- AJOUT MANUEL (NOUVEAU) ---
    if request.method == 'POST' and 'ajouter_manuel' in request.POST:
        try:
            base_mat = request.POST.get('matricule').strip()
            annee_obj = AnneeScolaire.objects.get(id=request.POST.get('annee'))
            pk_unique = f"{base_mat}-{annee_obj.libelle[:4]}"
            
            if Etudiant.objects.filter(matricule=pk_unique).exists():
                messages.error(request, "Cet étudiant existe déjà pour cette année.")
            else:
                etudiant = Etudiant.objects.create(
                    matricule=pk_unique,
                    nom=request.POST.get('nom').upper(),
                    prenom=request.POST.get('prenom').capitalize(),
                    classe_id=request.POST.get('classe'),
                    annee=annee_obj
                )
                DossierMemoire.objects.create(etudiant=etudiant)
                messages.success(request, f"L'étudiant {etudiant.nom} a été ajouté.")
        except Exception as e:
            messages.error(request, f"Erreur lors de l'ajout : {e}")
        return redirect('administration:liste_etudiants')

    if request.method == 'POST' and request.FILES.get('file_excel'):
        file = request.FILES['file_excel']
        try:
            # engine='openpyxl' est indispensable pour les fichiers .xlsx
            df = pd.read_excel(file, engine='openpyxl').fillna('')
            import_count, skip_count, error_count = 0, 0, 0
            
            for index, row in df.iterrows():
                try:
                    # Nettoyage du matricule (supprime les .0 si Excel le traite en nombre)
                    base_mat = str(row['matricule']).strip().split('.')[0]
                    
                    # Validation : Uniquement chiffres et au moins 4 caractères
                    if not base_mat.isdigit() or len(base_mat) < 4:
                        error_count += 1
                        continue

                    annee_libelle = str(row['annee_libelle']).strip()
                    classe_code = str(row['classe_code']).strip()
                    
                    # Identifiant unique composé pour la PK (Clé Primaire)
                    # Permet à un étudiant d'exister sur plusieurs années différentes
                    pk_unique = f"{base_mat}-{annee_libelle[:4]}"
                    
                    # Vérification si l'étudiant existe déjà pour cette année
                    if Etudiant.objects.filter(matricule=pk_unique).exists():
                        skip_count += 1
                        continue

                    # Récupération de la classe et de l'année (doivent exister en base)
                    classe_obj = Classe.objects.get(code=classe_code)
                    annee_obj = AnneeScolaire.objects.get(libelle=annee_libelle)

                    # 1. Création de l'étudiant
                    etudiant = Etudiant.objects.create(
                        matricule=pk_unique,
                        nom=str(row['nom']).strip().upper(),
                        prenom=str(row['prenom']).strip().capitalize(),
                        classe=classe_obj,
                        annee=annee_obj
                    )
                    
                    # 2. Initialisation du Dossier (Rôle futur du Chef de Dept)
                    # On crée l'objet DossierMemoire lié à l'étudiant
                    DossierMemoire.objects.create(etudiant=etudiant)
                    
                    import_count += 1

                except (Classe.DoesNotExist, AnneeScolaire.DoesNotExist):
                    error_count += 1
                    continue
                except Exception as e:
                    print(f"Erreur ligne {index}: {e} - views.py:1740")
                    error_count += 1
                    continue

            # Messages de retour personnalisés
            if import_count > 0:
                messages.success(request, f"{import_count} étudiants importés avec succès.")
            if skip_count > 0:
                messages.warning(request, f"{skip_count} matricules déjà existants pour cette année.")
            if error_count > 0:
                messages.error(request, f"{error_count} lignes ignorées (erreurs de format ou codes inexistants).")

        except Exception as e:
            messages.error(request, f"Impossible de lire le fichier : {str(e)}")
        
        return redirect('administration:liste_etudiants')

    # --- LOGIQUE DES FILTRES ET AFFICHAGE ---
    c_id = request.GET.get('classe')
    a_id = request.GET.get('annee')
    
    # On utilise select_related pour optimiser les performances (1 seule requête SQL)
    etudiants = Etudiant.objects.select_related('classe', 'annee').all()

    if c_id:
        etudiants = etudiants.filter(classe_id=c_id)
    if a_id:
        etudiants = etudiants.filter(annee_id=a_id)

    return render(request, 'administration/etudiants_liste.html', {
        'etudiants': etudiants,
        'classes': Classe.objects.all(),
        'annees': AnneeScolaire.objects.all()
    })